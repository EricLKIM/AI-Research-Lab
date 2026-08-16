"""Export weighted Analysis results in the user's selected format."""
from __future__ import annotations

import html as html_lib
import json
from datetime import datetime
from pathlib import Path

from research_lab.digest.topic_formatter import _slugify, TOPICS_DIR

EXTENSIONS = {"obsidian": "md", "markdown": "md", "text": "txt", "json": "json", "html": "html", "docx": "docx"}

HEADINGS = {
    "ko": {"title": "동향 분석", "summary": "전체 요약", "confirmed": "Confirmed Trend", "emerging": "Emerging Signal", "rumor": "Rumor", "contradictions": "상충 정보", "sources": "근거 자료"},
    "en": {"title": "Trend Analysis", "summary": "Overall Summary", "confirmed": "Confirmed Trend", "emerging": "Emerging Signal", "rumor": "Rumor", "contradictions": "Contradictions", "sources": "Evidence Sources"},
}

def _data(result: dict, topic: str, date: str) -> dict:
    return {"topic": topic, "date": date, "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M"), **result}

def _markdown(result: dict, topic: str, date: str, lang: str) -> str:
    h = HEADINGS.get(lang, HEADINGS["en"])
    lines = [f"# {h['title']}: {topic} ({date})", ""]
    if result.get("overall_summary"):
        lines += [f"## {h['summary']}", result["overall_summary"], ""]
    for key, title in (("confirmed_trends", h["confirmed"]), ("emerging_signals", h["emerging"]), ("rumors", h["rumor"])):
        items = result.get(key, [])
        lines.append(f"## {title}")
        if not items:
            lines.append("- None")
        for x in items:
            lines.append(f"### {x.get('title','')}")
            lines.append(f"- Confidence: {x.get('confidence',0)}% | Weight: {x.get('weight',0):.3f}")
            lines.append(f"- {x.get('summary','')}")
            if x.get("tags"):
                lines.append("- Tags: " + ", ".join(f"{t['tag']} ({t['confidence']}%)" for t in x["tags"]))
            if x.get("source_indices"):
                lines.append("- Sources: " + ", ".join(str(i) for i in x["source_indices"]))
            lines.append("")
    if result.get("contradictions"):
        lines += [f"## {h['contradictions']}"]
        for x in result["contradictions"]:
            lines.append(f"- **{x.get('topic','')}** — {x.get('summary','')} (sources: {x.get('source_indices',[])})")
        lines.append("")
    if result.get("cross_topic_evidence"):
        lines += ["## Cross-topic evidence"]
        for row in result["cross_topic_evidence"]:
            lines.append(
                f"- [{row.get('match_type','direct')}] {row.get('topic','')} / {row.get('tag','')}: "
                f"{row.get('summary','')} (confidence: {row.get('confidence', 0)}%)"
            )
        lines.append("")
    if result.get("time_series"):
        series = result["time_series"]
        lines += ["## Time-series trend", f"- Period: {series.get('period_start')} to {series.get('period_end')} ({series.get('period_days')} days)"]
        for signal in series.get("signals", []):
            if signal.get("direction") != "stable":
                lines.append(f"- {signal.get('tag')}: {signal.get('direction')} ({signal.get('before')} → {signal.get('recent')})")
        lines.append("")
    if result.get("data_quality"):
        quality = result["data_quality"]
        lines += ["## Data quality", f"- Snapshots: {quality.get('snapshot_count', 0)} | Articles: {quality.get('article_count', 0)} | Independent domains: {quality.get('independent_domain_count', 0)}", f"- Duplicate rate: {quality.get('duplicate_rate', 0):.1%} | Time-unknown rate: {quality.get('time_unknown_rate', 0):.1%}", f"- Coverage: {quality.get('coverage_note', 'unknown')}", ""]
    if result.get("alerts"):
        lines += ["## Alerts"]
        lines.extend(f"- {alert.get('message', '')}" for alert in result["alerts"])
        lines.append("")
    lines += [f"## {h['sources']}"]
    for i, s in enumerate(result.get("sources", [])):
        lines.append(f"### [{i}] {s.get('title','')}")
        lines.append(f"- Source: {s.get('source','')} | Kind: {s.get('kind','')}")
        lines.append(f"- Reliability: {s.get('reliability_score',0)}% | Freshness: {round(s.get('freshness_score',0)*100,1)}% | Weight: {s.get('weight',0):.3f}")
        if s.get("url"): lines.append(f"- URL: {s['url']}")
        lines.append("")
    return "\n".join(lines) + "\n"


def _obsidian_markdown(result: dict, topic: str, date: str, lang: str) -> str:
    """Give Analysis notes the same Obsidian-friendly metadata shape as research notes."""
    generated_at = datetime.now().astimezone().isoformat(timespec="seconds")
    frontmatter = [
        "---",
        f"date: {date}",
        f"topic: {json.dumps(topic, ensure_ascii=False)}",
        "tags:",
        "  - trend-analysis",
        *[f"  - {json.dumps(str(tag), ensure_ascii=False)}" for tag in result.get("analysis_tags", [])],
        f"generated_at: {json.dumps(generated_at)}",
        "---",
        "",
    ]
    return "\n".join(frontmatter) + _markdown(result, topic, date, lang)

def _text(result: dict, topic: str, date: str, lang: str) -> str:
    return _markdown(result, topic, date, lang).replace("# ", "").replace("## ", "[ ").replace("### ", "- ")

def _html(result: dict, topic: str, date: str, lang: str) -> str:
    md = _markdown(result, topic, date, lang)
    esc = html_lib.escape
    return f"<!doctype html><html><meta charset='utf-8'><title>{esc(topic)} Analysis</title><body><main><pre style='white-space:pre-wrap;font-family:Segoe UI,sans-serif'>{esc(md)}</pre></main></body></html>"

def save_analysis(result: dict, topic: str, date: str, vault_dir: Path, fmt: str = "obsidian", lang: str = "en") -> Path:
    fmt = (fmt or "obsidian").lower()
    if fmt not in EXTENSIONS: raise ValueError(f"Unsupported export format: {fmt}")
    out_dir = Path(vault_dir) / TOPICS_DIR / _slugify(topic)
    out_dir.mkdir(parents=True, exist_ok=True)
    # Analyses can be run several times in a day; retain each human-readable
    # result instead of overwriting the earlier one.
    timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
    path = out_dir / f"{timestamp}_analysis.{EXTENSIONS[fmt]}"
    if fmt == "json":
        path.write_text(json.dumps(_data(result, topic, date), ensure_ascii=False, indent=2), encoding="utf-8")
    elif fmt == "docx":
        try:
            from docx import Document
        except ImportError as e:
            raise RuntimeError("python-docx 패키지가 설치되어 있지 않습니다.") from e
        doc = Document(); doc.add_heading(f"Trend Analysis: {topic} ({date})", level=1)
        text = _markdown(result, topic, date, lang)
        for line in text.splitlines()[1:]:
            if line.startswith("### "): doc.add_heading(line[4:], level=3)
            elif line.startswith("## "): doc.add_heading(line[3:], level=2)
            elif line.strip(): doc.add_paragraph(line.lstrip("- "))
        doc.save(path)
    elif fmt == "html": path.write_text(_html(result, topic, date, lang), encoding="utf-8")
    elif fmt == "text": path.write_text(_text(result, topic, date, lang), encoding="utf-8")
    elif fmt == "obsidian": path.write_text(_obsidian_markdown(result, topic, date, lang), encoding="utf-8")
    else: path.write_text(_markdown(result, topic, date, lang), encoding="utf-8")
    return path
