"""
multi_format.py

다이제스트/주제 리서치 결과를 Obsidian 마크다운이 아닌 다른 형식으로도
내보낼 수 있게 해주는 모듈.

지원 형식:
- obsidian : 기존 방식 그대로 (YAML frontmatter, [[wikilink]], callout).
             DigestFormatter / TopicDigestFormatter를 그대로 사용한다.
- markdown : 옵시디언 전용 문법 없이 어디서나 통용되는 일반 마크다운(.md).
- text     : 마크다운 문법도 없는 순수 텍스트(.txt). 메모장 등에서 보기 좋음.
- json     : 분석 결과 전체를 구조화된 JSON으로 저장 (다른 프로그램에서 후처리할 때).
- html     : 브라우저에서 바로 열어볼 수 있는 간단한 HTML 페이지.
- docx     : Word 문서 (python-docx 필요).

폴더 구조(vault/digest/, vault/topics/{주제}/)는 형식과 무관하게 동일하게 유지되고,
파일 확장자만 형식에 맞게 바뀐다.

lang("ko"/"en")는 각 형식의 고정 틀 문구(섹션 제목 등)만 바꾼다. GPT가 작성한 실제
본문 내용의 언어는 analyzer 호출 시 넘긴 output_language로 이미 결정되어 있으므로,
lang은 보통 output_language를 resolve_ui_lang()으로 변환한 값을 그대로 넘겨주면 된다.
"""

from __future__ import annotations

import html as html_lib
import json
from dataclasses import asdict
from datetime import datetime
from pathlib import Path

from research_lab.analyzer.gpt import AnalysisResult
from research_lab.analyzer.topic_gpt import TopicAnalysisResult
from research_lab.digest.formatter import DIGEST_DIR, DigestFormatter
from research_lab.digest.topic_formatter import TOPICS_DIR, TopicDigestFormatter, _slugify
from research_lab.i18n import file_strings

EXPORT_FORMATS = ["obsidian", "markdown", "text", "json", "html", "docx"]
DEFAULT_FORMAT = "obsidian"

_EXT = {"obsidian": "md", "markdown": "md", "text": "txt", "json": "json", "html": "html", "docx": "docx"}

_HTML_STYLE = """<style>
body{font-family:-apple-system,'Segoe UI',sans-serif;background:#f3f6fa;color:#1b2733;margin:0}
main{max-width:760px;margin:0 auto;padding:32px 24px}
h1{color:#14508f} h2{color:#1d6fd6;border-bottom:2px solid #d7e0ea;padding-bottom:4px;margin-top:32px}
.meta{color:#5b6b7a;font-size:.9em}
.card{background:#fff;border:1px solid #d7e0ea;border-radius:10px;padding:14px 18px;margin:12px 0}
.source{color:#5b6b7a;font-size:.85em;margin:2px 0}
a{color:#0f9b8e}
</style>"""


def _new_docx_document():
    """python-docx를 지연 임포트한다 (설치 안 돼 있으면 안내 메시지와 함께 실패)."""
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx 패키지가 설치되어 있지 않습니다. `uv add python-docx` 실행 후 다시 시도하세요."
        ) from e
    return Document()


def _validate_format(fmt: str) -> str:
    fmt = (fmt or DEFAULT_FORMAT).strip().lower()
    if fmt not in EXPORT_FORMATS:
        raise ValueError(f"지원하지 않는 내보내기 형식: {fmt} (지원: {', '.join(EXPORT_FORMATS)})")
    return fmt


def _validate_lang(lang: str) -> str:
    return lang if lang in ("ko", "en") else "ko"


# ── AI 리서치 다이제스트 (AI Times + HuggingFace) ─────────────────────────

def save_digest(
    result: AnalysisResult, vault_dir: Path, fmt: str = DEFAULT_FORMAT, lang: str = "ko"
) -> Path:
    """분석 결과를 지정한 형식/언어로 vault_dir/digest/ 에 저장하고 저장된 경로를 반환한다."""
    fmt = _validate_format(fmt)
    lang = _validate_lang(lang)
    if fmt == "obsidian":
        return DigestFormatter(vault_dir, lang=lang).save(result)

    out_dir = Path(vault_dir) / DIGEST_DIR
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{result.date}.{_EXT[fmt]}"
    if fmt == "docx":
        _build_digest_docx(result, lang).save(path)
    else:
        path.write_text(_render_digest(result, fmt, lang), encoding="utf-8")
    return path


def _render_digest(r: AnalysisResult, fmt: str, lang: str) -> str:
    if fmt == "json":
        data = asdict(r)
        data["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        return json.dumps(data, ensure_ascii=False, indent=2)
    if fmt == "markdown":
        return _digest_markdown(r, lang)
    if fmt == "text":
        return _digest_text(r, lang)
    if fmt == "html":
        return _digest_html(r, lang)
    raise ValueError(f"지원하지 않는 형식: {fmt}")


def _digest_markdown(r: AnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [f"# {s['digest_title']} — {r.date}", "", f"_{s['generated_at']}: {now}_", ""]

    if r.trend_summary:
        lines.append(f"## {s['trend_summary_heading']}")
        lines += [f"{i}. {t}" for i, t in enumerate(r.trend_summary, 1)]
        lines.append("")

    if r.highlights:
        lines.append(f"## {s['highlights_heading']}")
        for h in r.highlights:
            lines.append(f"### {h.get('title', '')}")
            if h.get("url"):
                lines.append(f"- {s['link_label']}: {h['url']}")
            if h.get("source"):
                lines.append(f"- {s['source_label']}: {h['source']}")
            if h.get("why_important"):
                lines.append(f"- {h['why_important']}")
            lines.append("")

    if r.research_implications:
        lines.append(f"## {s['research_implications_heading']}")
        lines += [f"- {i}" for i in r.research_implications]
        lines.append("")

    if r.ip_perspective:
        lines.append(f"## {s['ip_perspective_heading']}")
        lines.append(r.ip_perspective)
        lines.append("")

    if r.suggested_nodes:
        lines.append(f"## {s['related_keywords_heading']}")
        for node in r.suggested_nodes:
            lines.append(f"- **{node.get('title', '')}**: {node.get('content', '')}")
        lines.append("")

    if r.error:
        lines.append(f"## {s['error_heading']}")
        lines.append(r.error)

    return "\n".join(lines) + "\n"


def _digest_text(r: AnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    lines = [f"{s['digest_title']} — {r.date}", f"({s['generated_at']}: {now})", "=" * 40, ""]

    if r.trend_summary:
        lines.append(f"[{s['trend_summary_heading']}]")
        lines += [f"{i}. {t}" for i, t in enumerate(r.trend_summary, 1)]
        lines.append("")

    if r.highlights:
        lines.append(f"[{s['highlights_heading']}]")
        for h in r.highlights:
            lines.append(f"- {h.get('title', '')} ({h.get('source', '')})")
            if h.get("url"):
                lines.append(f"  {h['url']}")
            if h.get("why_important"):
                lines.append(f"  {h['why_important']}")
        lines.append("")

    if r.research_implications:
        lines.append(f"[{s['research_implications_heading']}]")
        lines += [f"- {i}" for i in r.research_implications]
        lines.append("")

    if r.ip_perspective:
        lines.append(f"[{s['ip_perspective_heading']}]")
        lines.append(r.ip_perspective)
        lines.append("")

    return "\n".join(lines) + "\n"


def _digest_html(r: AnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    esc = html_lib.escape

    parts = [
        f"<!DOCTYPE html><html lang='{lang}'><head><meta charset='utf-8'>",
        f"<title>{esc(s['digest_title'])} — {esc(r.date)}</title>",
        _HTML_STYLE,
        "</head><body><main>",
        f"<h1>🤖 {esc(s['digest_title'])} — {esc(r.date)}</h1>",
        f"<p class='meta'>{esc(s['generated_at'])}: {esc(now)}</p>",
    ]

    if r.trend_summary:
        parts.append(f"<h2>{esc(s['trend_summary_heading'])}</h2><ol>")
        parts += [f"<li>{esc(t)}</li>" for t in r.trend_summary]
        parts.append("</ol>")

    if r.highlights:
        parts.append(f"<h2>{esc(s['highlights_heading'])}</h2>")
        for h in r.highlights:
            title, url = esc(h.get("title", "")), h.get("url", "")
            title_html = f"<a href='{esc(url)}'>{title}</a>" if url else title
            parts.append("<div class='card'>")
            parts.append(f"<h3>{title_html}</h3>")
            if h.get("source"):
                parts.append(f"<p class='source'>{esc(s['source_label'])}: {esc(h['source'])}</p>")
            if h.get("why_important"):
                parts.append(f"<p>{esc(h['why_important'])}</p>")
            parts.append("</div>")

    if r.research_implications:
        parts.append(f"<h2>{esc(s['research_implications_heading'])}</h2><ul>")
        parts += [f"<li>{esc(i)}</li>" for i in r.research_implications]
        parts.append("</ul>")

    if r.ip_perspective:
        parts.append(f"<h2>{esc(s['ip_perspective_heading'])}</h2>")
        parts.append(f"<p>{esc(r.ip_perspective)}</p>")

    parts.append("</main></body></html>")
    return "\n".join(parts)


def _build_digest_docx(r: AnalysisResult, lang: str):
    """AI 다이제스트 분석 결과를 Word(.docx) 문서로 만든다."""
    s = file_strings(lang)
    doc = _new_docx_document()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")

    doc.add_heading(f"{s['digest_title']} — {r.date}", level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"{s['generated_at']}: {now}").italic = True

    if r.trend_summary:
        doc.add_heading(s["trend_summary_heading"], level=2)
        for t in r.trend_summary:
            doc.add_paragraph(t, style="List Number")

    if r.highlights:
        doc.add_heading(s["highlights_heading"], level=2)
        for h in r.highlights:
            doc.add_heading(h.get("title", ""), level=3)
            if h.get("source"):
                doc.add_paragraph(f"{s['source_label']}: {h['source']}")
            if h.get("url"):
                doc.add_paragraph(h["url"])
            if h.get("why_important"):
                doc.add_paragraph(h["why_important"])

    if r.research_implications:
        doc.add_heading(s["research_implications_heading"], level=2)
        for i in r.research_implications:
            doc.add_paragraph(i, style="List Bullet")

    if r.ip_perspective:
        doc.add_heading(s["ip_perspective_heading"], level=2)
        doc.add_paragraph(r.ip_perspective)

    if r.suggested_nodes:
        doc.add_heading(s["related_keywords_heading"], level=2)
        for node in r.suggested_nodes:
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(node.get("title", "")).bold = True
            p.add_run(f": {node.get('content', '')}")

    if r.error:
        doc.add_heading(s["error_heading"], level=2)
        doc.add_paragraph(r.error)

    return doc


# ── 주제 리서치 ─────────────────────────────────────────────────────────

def save_topic_digest(
    result: TopicAnalysisResult, vault_dir: Path, fmt: str = DEFAULT_FORMAT, lang: str = "ko"
) -> Path:
    """주제 분석 결과를 지정한 형식/언어로 vault_dir/topics/{주제}/ 에 저장한다."""
    fmt = _validate_format(fmt)
    lang = _validate_lang(lang)
    if fmt == "obsidian":
        return TopicDigestFormatter(vault_dir, lang=lang).save(result)

    topic_slug = _slugify(result.topic)
    out_dir = Path(vault_dir) / TOPICS_DIR / topic_slug
    out_dir.mkdir(parents=True, exist_ok=True)
    path = out_dir / f"{result.date}.{_EXT[fmt]}"
    if fmt == "docx":
        _build_topic_docx(result, lang).save(path)
    else:
        path.write_text(_render_topic(result, fmt, lang), encoding="utf-8")
    return path


def _render_topic(r: TopicAnalysisResult, fmt: str, lang: str) -> str:
    if fmt == "json":
        data = asdict(r)
        data["generated_at"] = datetime.now().strftime("%Y-%m-%d %H:%M")
        return json.dumps(data, ensure_ascii=False, indent=2)
    if fmt == "markdown":
        return _topic_markdown(r, lang)
    if fmt == "text":
        return _topic_text(r, lang)
    if fmt == "html":
        return _topic_html(r, lang)
    raise ValueError(f"지원하지 않는 형식: {fmt}")


def _topic_markdown(r: TopicAnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    title = s["topic_title_fmt"].format(topic=r.topic)
    lines = [f"# {title} ({r.date})", "", f"_{s['generated_at']}: {now}_", ""]

    if r.trend_summary:
        lines.append(f"## {s['topic_trend_heading']}")
        lines += [f"- {t}" for t in r.trend_summary]
        lines.append("")

    if r.highlights:
        lines.append(f"## {s['topic_highlights_heading']}")
        for h in r.highlights:
            lines.append(f"### {h.get('title', '')}")
            if h.get("url"):
                lines.append(f"- {s['link_label']}: {h['url']}")
            if h.get("source"):
                lines.append(f"- {s['source_label']}: {h['source']}")
            if h.get("why_important"):
                lines.append(f"- {h['why_important']}")
            lines.append("")

    if r.collected_articles:
        lines.append(f"## {s['topic_sources_heading']} ({len(r.collected_articles)})")
        for article in r.collected_articles:
            title = str(article.get("title") or "Untitled source")
            url = str(article.get("url") or "")
            source = str(article.get("source") or "")
            platform = str(article.get("platform") or "")
            kind = str(article.get("kind") or "")
            date = str(article.get("date") or "")
            details = " · ".join(value for value in (source, platform, kind, date) if value)
            link = f"[{title}]({url})" if url else title
            lines.append(f"- {link}" + (f" — {details}" if details else ""))
        lines.append("")

    if r.key_takeaways:
        lines.append(f"## {s['topic_takeaways_heading']}")
        lines += [f"- {t}" for t in r.key_takeaways]
        lines.append("")

    return "\n".join(lines) + "\n"


def _topic_text(r: TopicAnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    title = s["topic_title_fmt"].format(topic=r.topic)
    lines = [f"{title} ({r.date})", f"({s['generated_at']}: {now})", "=" * 40, ""]

    if r.trend_summary:
        lines.append(f"[{s['topic_trend_heading']}]")
        lines += [f"- {t}" for t in r.trend_summary]
        lines.append("")

    if r.highlights:
        lines.append(f"[{s['topic_highlights_heading']}]")
        for h in r.highlights:
            lines.append(f"- {h.get('title', '')} ({h.get('source', '')})")
            if h.get("url"):
                lines.append(f"  {h['url']}")
            if h.get("why_important"):
                lines.append(f"  {h['why_important']}")
        lines.append("")

    if r.key_takeaways:
        lines.append(f"[{s['topic_takeaways_heading']}]")
        lines += [f"- {t}" for t in r.key_takeaways]
        lines.append("")

    return "\n".join(lines) + "\n"


def _topic_html(r: TopicAnalysisResult, lang: str) -> str:
    s = file_strings(lang)
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    esc = html_lib.escape
    title = s["topic_title_fmt"].format(topic=r.topic)
    parts = [
        f"<!DOCTYPE html><html lang='{lang}'><head><meta charset='utf-8'>",
        f"<title>{esc(title)} — {esc(r.date)}</title>",
        _HTML_STYLE,
        "</head><body><main>",
        f"<h1>🔎 {esc(title)} — {esc(r.date)}</h1>",
        f"<p class='meta'>{esc(s['generated_at'])}: {esc(now)}</p>",
    ]

    if r.trend_summary:
        parts.append(f"<h2>{esc(s['topic_trend_heading'])}</h2><ul>")
        parts += [f"<li>{esc(t)}</li>" for t in r.trend_summary]
        parts.append("</ul>")

    if r.highlights:
        parts.append(f"<h2>{esc(s['topic_highlights_heading'])}</h2>")
        for h in r.highlights:
            h_title, url = esc(h.get("title", "")), h.get("url", "")
            title_html = f"<a href='{esc(url)}'>{h_title}</a>" if url else h_title
            parts.append("<div class='card'>")
            parts.append(f"<h3>{title_html}</h3>")
            if h.get("source"):
                parts.append(f"<p class='source'>{esc(s['source_label'])}: {esc(h['source'])}</p>")
            if h.get("why_important"):
                parts.append(f"<p>{esc(h['why_important'])}</p>")
            parts.append("</div>")

    if r.key_takeaways:
        parts.append(f"<h2>{esc(s['topic_takeaways_heading'])}</h2><ul>")
        parts += [f"<li>{esc(t)}</li>" for t in r.key_takeaways]
        parts.append("</ul>")

    parts.append("</main></body></html>")
    return "\n".join(parts)


def _build_topic_docx(r: TopicAnalysisResult, lang: str):
    """주제 리서치 분석 결과를 Word(.docx) 문서로 만든다."""
    s = file_strings(lang)
    doc = _new_docx_document()
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    title = s["topic_title_fmt"].format(topic=r.topic)

    doc.add_heading(f"{title} ({r.date})", level=1)
    meta = doc.add_paragraph()
    meta.add_run(f"{s['generated_at']}: {now}").italic = True

    if r.trend_summary:
        doc.add_heading(s["topic_trend_heading"], level=2)
        for t in r.trend_summary:
            doc.add_paragraph(t, style="List Bullet")

    if r.highlights:
        doc.add_heading(s["topic_highlights_heading"], level=2)
        for h in r.highlights:
            doc.add_heading(h.get("title", ""), level=3)
            if h.get("source"):
                doc.add_paragraph(f"{s['source_label']}: {h['source']}")
            if h.get("url"):
                doc.add_paragraph(h["url"])
            if h.get("why_important"):
                doc.add_paragraph(h["why_important"])

    if r.key_takeaways:
        doc.add_heading(s["topic_takeaways_heading"], level=2)
        for t in r.key_takeaways:
            doc.add_paragraph(t, style="List Bullet")

    return doc
